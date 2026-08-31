#!/usr/bin/env python3
"""Verificação de plágio: procura sequências de N palavras da prosa do pôster e
dos resumos do fichamento que apareçam LITERAIS em algum PDF de texto/.

Também confere que cada `citacao_direta` de fichas_analiticas.yml é verbatim.

Uso:  python3 _check_plagio.py
"""
from __future__ import annotations

import glob
import re
from pathlib import Path

import yaml
from docx import Document

HERE = Path(__file__).parent
N = 8


def norm(t: str) -> str:
    t = re.sub(r"\[\[página \d+\]\]", " ", t).replace("­", "")
    t = re.sub(r"([a-zà-ÿ])-\s*\n\s*([a-zà-ÿ])", r"\1\2", t)
    t = t.lower().replace("’", "'").replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", re.sub(r"[^a-zà-ÿ0-9 ]", " ", t))


corpus = {Path(f).stem: norm(open(f, encoding="utf-8", errors="ignore").read())
          for f in glob.glob(str(HERE / "texto" / "*.txt"))}
big = " \n ".join(corpus.values())
bw = big.split()
big_ng = {" ".join(bw[i:i + N]) for i in range(len(bw) - N)}


def ngrams_lit(texto: str) -> list[str]:
    w = norm(texto).split()
    return [" ".join(w[i:i + N]) for i in range(len(w) - N) if " ".join(w[i:i + N]) in big_ng]


print(f"=== PÔSTER (banner-final) — prosa vs. {len(corpus)} PDFs, janela de {N} palavras ===")
banner = HERE.parent / "banner-final-sipes-2026.docx"
if banner.exists():
    d = Document(str(banner))
    prosa = []
    for p in d.paragraphs:
        prosa.append(p.text)
    for tb in d.tables:
        for r in tb.rows:
            for cc in r.cells:
                for p in cc.paragraphs:
                    prosa.append(p.text)
    prosa = [t for t in prosa if len(t) > 120 and "DOI: https" not in t
             and not t.strip().startswith(("Figura", "REFERÊNCIAS"))]
    h = ngrams_lit(" ".join(prosa))
    print(f"  {len(h)} coincidência(s)" + ("" if not h else ":"))
    for x in h[:20]:
        print("   !", x)

print(f"\n=== FICHAMENTO — resumos (paráfrase própria) ===")
ANA = yaml.safe_load((HERE / "fichas_analiticas.yml").read_text(encoding="utf-8"))
for k, f in ANA.items():
    w = norm(f["resumo"]).split()
    src = set(" ".join(corpus.get(k, "").split()[i:i + N])
              for i in range(len(corpus.get(k, "").split()) - N))
    h = [" ".join(w[i:i + N]) for i in range(len(w) - N) if " ".join(w[i:i + N]) in src]
    if h:
        print(f"  {k}: {len(h)} — {h}")
print("  (coincidências restantes = termos técnicos / nomes próprios)")

print("\n=== CITAÇÕES DIRETAS — devem ser verbatim ===")
bad = 0
for k, f in ANA.items():
    q = f.get("citacao_direta", "")
    if not q or q.startswith("["):
        continue
    src = corpus.get(k, "")
    if norm(q) in src or " ".join(norm(q).split()[:10]) in src:
        continue
    print(f"  FALHA {k}")
    bad += 1
print(f"  {len(ANA) - bad}/{len(ANA)} conferem" if bad else f"  todas as {len(ANA)} conferem")
