#!/usr/bin/env python3
"""Gera docs/sipes-2026/banner-sipes-2026-modelo.docx — modelo de pôster/banner
para o V SIPES 2026, com a logomarca do evento (obrigatória).

O banner do evento é 120 x 90 cm. O Word limita a página a ~55,9 cm, então o
modelo vem em ESCALA 1:3 (30 x 40 cm, proporção 3:4). Ao finalizar: definir a
página como 90 cm (largura) x 120 cm (altura) OU exportar em PDF e ampliar 3x;
manter 300 dpi nas imagens.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

AQUA = "1C8CB0"
AQUA_RGB = RGBColor(0x1C, 0x8C, 0xB0)
AQUA_LIGHT = "E8F4F8"
NAVY = RGBColor(0x12, 0x30, 0x3F)
GREY = RGBColor(0x60, 0x68, 0x75)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
PLACEHOLDER = RGBColor(0x9A, 0xA3, 0xB0)

HERE = Path(__file__).parent
LOGO = HERE / "assets" / "sipes-logo.png"
DIAG = HERE / "assets" / "diagrama-ieas.png"
OUT = HERE / "banner-sipes-2026-modelo.docx"


def shade(el, fill_hex):
    """Aplica cor de fundo a uma célula ou parágrafo (elemento _tc ou _p)."""
    pr = el._tc.get_or_add_tcPr() if hasattr(el, "_tc") else el._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pr.append(shd)


def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        e.set(qn("w:color"), "auto")
        borders.append(e)
    tblPr.append(borders)


def run(par, text, *, size, bold=False, color=None, italic=False):
    r = par.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return r


def para(cell_or_doc, text="", *, size=10, bold=False, color=None, italic=False,
         align=None, space_after=4, space_before=0):
    par = cell_or_doc.add_paragraph()
    if align is not None:
        par.alignment = align
    par.paragraph_format.space_after = Pt(space_after)
    par.paragraph_format.space_before = Pt(space_before)
    par.paragraph_format.line_spacing = 1.1
    if text:
        run(par, text, size=size, bold=bold, color=color, italic=italic)
    return par


def section_title(cell, text):
    p = para(cell, text, size=12.5, bold=True, color=WHITE, space_before=6, space_after=3)
    shade(p, AQUA)
    p.paragraph_format.left_indent = Pt(4)
    return p


def body(cell, text, *, italic=False, size=9.5):
    return para(cell, text, size=size, color=NAVY, italic=italic, align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def ph(cell, text):
    """Texto de instrução (placeholder) em cinza itálico."""
    return para(cell, text, size=9, color=PLACEHOLDER, italic=True)


# ── documento ────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(30), Cm(40)          # escala 1:3 de 90 x 120
sec.left_margin = sec.right_margin = Cm(1.4)
sec.top_margin = sec.bottom_margin = Cm(1.2)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10)

# nota de escala (fora da área visual do pôster, no topo)
note = para(doc, "MODELO EM ESCALA 1:3 — pôster final: 90 cm (largura) × 120 cm "
                 "(altura), retrato. Ao finalizar, redefinir o tamanho da página "
                 "ou exportar em PDF e ampliar 3×; imagens a 300 dpi. A logomarca "
                 "do V SIPES no cabeçalho é obrigatória.",
            size=7.5, color=GREY, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=6)

# ── CABEÇALHO ────────────────────────────────────────────────────────
head = doc.add_table(rows=1, cols=2)
head.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(head)
head.columns[0].width = Cm(9.5)
head.columns[1].width = Cm(17.7)
hc0, hc1 = head.rows[0].cells
shade(hc0, AQUA_LIGHT)
shade(hc1, AQUA_LIGHT)
if LOGO.exists():
    lp = hc0.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(str(LOGO), width=Cm(8.6))
para(hc1, "V Simpósio Internacional de Pesquisa em Estilos de Vida & Saúde",
     size=11, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
para(hc1, "Recife-PE · 25 a 27 de novembro de 2026", size=9.5, color=GREY,
     align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
para(hc1, "Modalidade: Pôster   ·   Eixo temático: [selecionar]", size=9,
     color=AQUA_RGB, align=WD_ALIGN_PARAGRAPH.RIGHT)

para(doc, "", space_after=2)

# ── TÍTULO + AUTORES ─────────────────────────────────────────────────
para(doc, "[ TÍTULO DO TRABALHO — no máximo 50 palavras, só a primeira letra "
          "maiúscula ]", size=19, bold=True, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
para(doc, "[ Sobrenome, N.¹  ·  Sobrenome, N.²  ·  … (até 6 autores; o 1º "
          "apresenta) ]", size=11, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_after=1)
para(doc, "[ ¹Instituição, cidade/UF  ·  ²Instituição  ·  e-mail do 1º autor ]",
     size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

# ── CORPO EM 2 COLUNAS ───────────────────────────────────────────────
body_t = doc.add_table(rows=1, cols=2)
body_t.alignment = WD_TABLE_ALIGNMENT.CENTER
no_borders(body_t)
body_t.columns[0].width = Cm(13.4)
body_t.columns[1].width = Cm(13.4)
L, R = body_t.rows[0].cells
for c in (L, R):
    c.paragraphs[0].paragraph_format.space_after = Pt(0)

# coluna esquerda
section_title(L, "INTRODUÇÃO")
body(L, "Alocação de recursos em saúde e carga de doença não se distribuem de "
        "forma equivalente no território. Os portais de transparência mostram "
        "quanto se gasta, não se o gasto acompanha a necessidade de cada "
        "município.")
ph(L, "[ 2–3 frases com 1 referência: iniquidade territorial em saúde / "
      "financiamento do SUS / determinação social pelo saneamento ]")

section_title(L, "OBJETIVOS")
body(L, "Desenvolver e disponibilizar o Farol da Saúde & Saneamento (Farol-SS), "
        "monitor territorial de código aberto que quantifica o alinhamento entre "
        "necessidade sanitária e alocação de recursos nos 185 municípios de "
        "Pernambuco (2020–2024).")
ph(L, "[ acrescentar 2 objetivos específicos ]")

section_title(L, "METODOLOGIA")
body(L, "Pipeline reprodutível (camadas bronze–silver–gold; DuckDB sobre "
        "Parquet; proveniência com SHA-256) integrou oito fontes federais "
        "abertas num grão único de município × ano: SINAN, internações do "
        "SIH-SUS por DRSAI, Censo 2022 (saneamento), CadÚnico (extrema pobreza), "
        "SIOPS (L2), PNCP e Compras.gov.br (L3), Portal da Transparência (L1).")
body(L, "IEAS — Índice de Efetividade da Alocação Sanitária: dois eixos, "
        "Necessidade (epidemiológico 0,40 + saneamento 0,35 + vulnerabilidade "
        "0,25) e Alocação (L1 0,35 + L2 0,40 + L3 0,25, R$/hab deflacionado), "
        "cada um convertido em ranque percentil ∈ [0,1] no estado. "
        "gap = ranque(A) − ranque(N) colore um semáforo de quatro cores; "
        "\"regra do cinza\" não publica índice sem cobertura mínima de dados. "
        "Quatro detectores geram alertas explicáveis em linguagem natural.")
if DIAG.exists():
    fp = L.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(4)
    fp.add_run().add_picture(str(DIAG), width=Cm(10.5))
    para(L, "Figura 1. Construção do IEAS — das oito fontes ao semáforo.",
         size=8, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

# coluna direita
section_title(R, "RESULTADOS")
body(R, "O IEAS foi calculado para 921 dos 925 município-anos (185/185 em 2024). "
        "Distribuição do farol (todos os anos): 383 alinhado, 202 necessidade "
        "não atendida, 176 alocação acima da necessidade, 160 subalocação leve, "
        "4 sem dado. Em 2024: 104 verde, 60 azul, 19 amarelo, 2 vermelho.")
body(R, "777 alertas emitidos: 570 suspeita de desabastecimento (concentrados "
        "nos anos de surto de arbovirose), 202 desalinhamento estrutural, "
        "4 alocação abaixo do esperado, 1 suspeita de sobrepreço.")
body(R, "Coberturas censitárias medianas em PE: água 64%, esgoto 48%, coleta de "
        "lixo 74% — necessidade elevada em quase todo o estado.")
fig2 = R.add_paragraph()
fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fig2.add_run("[  INSERIR AQUI: mapa coroplético do farol de PE — 2024  ]\n"
                  "[  exportar da página Farol do painel, PNG 300 dpi  ]")
fr.font.size = Pt(9)
fr.font.color.rgb = PLACEHOLDER
fr.italic = True
shade(fig2, "F1F3F6")
para(R, "Figura 2. Farol da alocação sanitária — Pernambuco, 2024.", size=8,
     color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
ph(R, "[ opcional: caixa com 1 exemplo de alerta em linguagem natural, ex.: "
      "\"Incidência de dengue no percentil 75+ de PE em 2024, sem contratação "
      "de larvicida/inseticida no PNCP para o município.\" ]")

section_title(R, "CONCLUSÕES")
body(R, "É viável, a partir exclusivamente de dados abertos, produzir um "
        "instrumento reprodutível e explicável de apoio ao controle social e à "
        "priorização de auditoria e de investimento em saúde e saneamento. O "
        "método é transferível a outras unidades da federação.")

section_title(R, "LIMITAÇÕES")
body(R, "L1 é um proxy (transferências sociais, não repasse setorial de saúde); "
        "saneamento é um retrato de 2022; a adesão ao PNCP cresce ano a ano "
        "(coluna de maturidade indica a confiança por ano).")

para(doc, "", space_after=2)

# ── RODAPÉ ───────────────────────────────────────────────────────────
foot = doc.add_table(rows=1, cols=2)
no_borders(foot)
foot.columns[0].width = Cm(20.2)
foot.columns[1].width = Cm(7.0)
fc0, fc1 = foot.rows[0].cells
shade(fc0, AQUA_LIGHT)
shade(fc1, AQUA_LIGHT)
para(fc0, "REFERÊNCIAS", size=9, bold=True, color=NAVY, space_after=1)
ph(fc0, "[ 5 a 8 referências, ABNT NBR 6023, fonte 7–8 pt — só as efetivamente "
        "citadas no pôster ]")
para(fc1, "Painel: farol-ss.streamlit.app", size=8, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
para(fc1, "[ QR code do painel ]", size=8, color=PLACEHOLDER, italic=True,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
para(fc1, "Código: github.com/protazoarium/farol-ss", size=7, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER)

para(doc, "V SIPES 2026 · Recife-PE · 25–27 nov · Dados sob domínio público das "
          "fontes federais (IBGE, DATASUS, PNCP, CGU, MDS)", size=7.5, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=4)

doc.save(str(OUT))
print("ok:", OUT, OUT.stat().st_size, "bytes")
