#!/usr/bin/env python3
"""Gera docs/sipes-2026/projeto-sipes-2026.docx — o projeto para construção do
trabalho a ser submetido ao V SIPES (Recife-PE, 25–27/nov/2026)."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

AQUA = RGBColor(0x1C, 0x8C, 0xB0)
NAVY = RGBColor(0x12, 0x30, 0x3F)
GREY = RGBColor(0x54, 0x5C, 0x6B)

HERE = Path(__file__).parent
LOGO = HERE / "assets" / "sipes-logo.png"
DIAG = HERE / "assets" / "diagrama-ieas.png"
OUT = HERE / "projeto-sipes-2026.docx"


def base_style(doc: Document) -> None:
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15
    for name, size in (("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        h = doc.styles[name]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.color.rgb = NAVY
        h.font.bold = True
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)


def p(doc, text="", *, bold=False, italic=False, size=11, color=None, align=None, style=None):
    par = doc.add_paragraph(style=style)
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return par


def bullet(doc, text, level=0):
    par = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    par.paragraph_format.space_after = Pt(2)
    return par


def kv_table(doc, rows, widths=(4.2, 12.3)):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = k
        cells[1].text = v
        cells[0].paragraphs[0].runs[0].bold = True
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def grid_table(doc, header, rows, widths):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(9.5)
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            for par in cells[i].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(9.5)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


# ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Cm(2.2))
base_style(doc)

# capa
if LOGO.exists():
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run().add_picture(str(LOGO), width=Cm(9))
p(doc, "Projeto para construção do trabalho", bold=True, size=20, color=NAVY,
  align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, "V Simpósio Internacional de Pesquisa em Estilos de Vida & Saúde — SIPES 2026",
  size=12, color=AQUA, align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, "Recife-PE · 25 a 27 de novembro de 2026 (em conjunto com SAIA e CIISE)",
  size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
p(doc, "Objeto: Farol da Saúde & Saneamento (Farol-SS / Painel-SS) — monitor "
       "territorial de efetividade da alocação sanitária para os 185 municípios "
       "de Pernambuco, construído exclusivamente com dados abertos federais.",
  italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
p(doc, "Documento de planejamento — não é o texto final de submissão. Uso interno "
       "da equipe autora.", size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# 1
doc.add_heading("1. Identificação do trabalho", 1)
kv_table(doc, [
    ("Evento", "V SIPES 2026 — Simpósio Internacional de Pesquisa em Estilos de Vida & Saúde"),
    ("Local / data", "Recife-PE · 25–27 de novembro de 2026"),
    ("Modalidade pretendida", "Apresentação em pôster (banner 120 × 90 cm) — com opção "
     "de indicar interesse em apresentação oral (10 min + 5 min de arguição)"),
    ("Área temática (a confirmar no formulário)", "Inovação tecnológica em saúde / "
     "Sistemas de informação e gestão em saúde — entre os 11 eixos do evento"),
    ("Título provisório", "Farol da Saúde & Saneamento: um índice territorial de "
     "efetividade da alocação sanitária construído com dados abertos para os 185 "
     "municípios de Pernambuco (2020–2024)"),
    ("Título alternativo (mais curto)", "Alinhamento entre necessidade e gasto em "
     "saúde nos municípios de Pernambuco: o Índice de Efetividade da Alocação "
     "Sanitária (IEAS)"),
    ("Autores (até 6; 1º autor = responsável e apresentador)",
     "[nome completo, e-mail, instituição] — preencher; o 1º autor só pode ser "
     "principal em um trabalho"),
    ("Palavras-chave (3, confirmadas no DeCS)",
     "Saneamento [D012499]; Alocação de Recursos [D040841]; "
     "Gastos em Saúde [D005102]"),
])
p(doc, "Regras do evento que condicionam o texto (fonte: sipes.com.br/trabalhos, "
       "consultado em 30/08/2026 — reconferir antes de enviar):", bold=True)
bullet(doc, "Resumo em parágrafo único, sem divisão de seções, com no máximo "
            "2.500 caracteres com espaços, contendo introdução, objetivos, "
            "metodologia, resultados e conclusões.")
bullet(doc, "Título com no máximo 50 palavras, apenas a primeira letra maiúscula.")
bullet(doc, "Três palavras-chave dos vocabulários DeCS ou MeSH.")
bullet(doc, "Até 6 autores; submissão só pelo formulário on-line; trabalho inédito "
            "(não apresentado em outro evento); qualidade textual é critério de "
            "avaliação; não há edição de autoria após o aceite.")
bullet(doc, "Pôster: 120 × 90 cm, com a logomarca do evento obrigatória; design "
            "livre; apresentação em português ou inglês.")

# 2
doc.add_heading("2. Contextualização — o que é o Farol-SS", 1)
p(doc, "Portais de transparência informam quanto se gasta em saúde, mas não se "
       "esse gasto acompanha a necessidade de cada território. O Farol-SS é um "
       "monitor de código aberto que torna esse descompasso mensurável para os "
       "185 municípios de Pernambuco, no período 2020–2024, cruzando oito fontes "
       "federais abertas num grão único de município × ano e produzindo o "
       "Índice de Efetividade da Alocação Sanitária (IEAS).")
p(doc, "Entregas já concluídas (base empírica do trabalho):", bold=True)
bullet(doc, "Pipeline de dados reprodutível (arquitetura em camadas bronze–silver–"
            "gold; DuckDB sobre Parquet; proveniência com SHA-256 em manifest.json).")
bullet(doc, "IEAS calculado para 921 dos 925 município-anos (185/185 em 2024).")
bullet(doc, "Quatro detectores de anomalia com alertas explicáveis em linguagem "
            "natural (777 alertas emitidos).")
bullet(doc, "Painel web público (Streamlit), API aberta em JSON/CSV (FastAPI) e "
            "relatório técnico completo.")
bullet(doc, "Submetido ao 2º Concurso de Reúso de Dados Abertos da CGU.")
p(doc, "Painel: https://farol-ss.streamlit.app  ·  Código: "
       "https://github.com/protazoarium/farol-ss  ·  Relatório técnico: "
       "docs/relatorio-tecnico.md do repositório.", size=9, color=GREY)

# 3
doc.add_heading("3. Objetivo do trabalho no SIPES", 1)
p(doc, "Objetivo geral: apresentar o Farol-SS como demonstração de que é possível, "
       "a partir exclusivamente de dados abertos, produzir um instrumento "
       "reprodutível e explicável de apoio ao controle social e à priorização de "
       "auditoria e de investimento em saúde e saneamento.")
p(doc, "Objetivos específicos (recorte para resumo + pôster):", bold=True)
bullet(doc, "Descrever a construção do IEAS (fontes, eixos, pesos, normalização "
            "por ranque percentil, regra do cinza).")
bullet(doc, "Apresentar o panorama estadual 2020–2024 (distribuição do farol; "
            "municípios em necessidade não atendida).")
bullet(doc, "Ilustrar os alertas explicáveis como ferramenta de vigilância e "
            "controle social (desabastecimento de insumos, sobrepreço, "
            "desalinhamento estrutural).")
bullet(doc, "Discutir limitações (proxy de L1; saneamento como retrato de 2022; "
            "maturidade do PNCP) e transferibilidade a outras UFs.")

# 4
doc.add_heading("4. Fundamentação teórica — protocolo de revisão da literatura", 1)
p(doc, "A fundamentação será construída por extração sistematizada de artigos → "
       "fichamento estruturado → redação a partir dos fichamentos. Meta mínima: "
       "12–18 referências efetivamente lidas e fichadas, das quais 8–12 citadas "
       "no resumo/pôster e no eventual texto expandido.")

doc.add_heading("4.1 Bases e estratégia de busca", 2)
grid_table(doc,
    ["Base", "Uso", "Observação"],
    [
     ["PubMed / MEDLINE", "literatura internacional (equidade, financiamento, DRSAI)", "termos MeSH"],
     ["BVS / LILACS", "literatura latino-americana e brasileira", "termos DeCS"],
     ["SciELO", "periódicos brasileiros de saúde coletiva", "Cad. Saúde Pública, Ciênc. Saúde Colet., Rev. Saúde Pública"],
     ["Scopus / Web of Science", "indicadores compostos, open government data", "se houver acesso institucional"],
     ["Google Scholar", "literatura cinzenta, teses, relatórios (Funasa, IPEA, OPAS, Ipea)", "rastrear citações"],
     ["Repositórios de dados/gov", "documentação metodológica de SINAN, SIH, SIOPS, PNCP, Censo", "não conta como referência teórica"],
    ],
    [3.2, 6.5, 6.5])
p(doc, "Descritores da submissão (confirmados no DeCS): Saneamento [D012499] · "
       "Alocação de Recursos [D040841] · Gastos em Saúde [D005102]. As strings "
       "de busca abaixo são mais amplas de propósito — servem para achar a "
       "literatura, não para o campo de palavras-chave do formulário.", bold=True)
grid_table(doc,
    ["Eixo temático", "Descritores / termos (DeCS · MeSH · livres)"],
    [
     ["Alocação de recursos e equidade em saúde",
      "\"Alocação de Recursos\" [D040841]/\"Resource Allocation\"; \"Equidade "
      "em Saúde\"/\"Health Equity\"; \"needs-based funding\"; \"territorial "
      "justice\" AND health"],
     ["Gastos em saúde e financiamento do SUS",
      "\"Gastos em Saúde\" [D005102]/\"Health Expenditures\"; \"Financiamento "
      "da Assistência à Saúde\"; \"Sistema Único de Saúde\" AND (financiamento "
      "OR subfinanciamento OR \"per capita\")"],
     ["Saneamento e saúde / DRSAI",
      "\"Saneamento\" [D012499]/\"Sanitation\"; \"doenças relacionadas ao "
      "saneamento ambiental inadequado\"; \"waterborne diseases\"; \"Saneamento\" "
      "AND (morbidade OR internação OR hospitalização)"],
     ["Vigilância e sistemas de informação em saúde",
      "\"Sistemas de Informação em Saúde\"/\"Health Information Systems\"; "
      "\"Vigilância em Saúde Pública\"; \"SINAN\"; \"record linkage\" AND health"],
     ["Dados abertos governamentais e reúso",
      "\"open government data\"; \"open data\" AND (reuse OR \"public value\" OR "
      "accountability); \"dados abertos\" AND (saúde OR governo)"],
     ["Indicadores compostos / índices territoriais",
      "\"composite indicator\"; \"index construction\" AND health; \"percentile "
      "rank\" normalization; \"small-area\" AND (health OR deprivation) AND index"],
     ["Alertas / red flags em contratações públicas",
      "\"corruption risk\" AND \"public procurement\"; \"red flags\" procurement; "
      "\"price analysis\" AND (medicines OR \"health supplies\") AND public"],
    ],
    [4.3, 12.0])

doc.add_heading("4.2 Critérios de inclusão e exclusão", 2)
bullet(doc, "Inclusão: artigos revisados por pares, relatórios técnicos de órgãos "
            "oficiais (OPAS, Funasa, IPEA, Ministério da Saúde, Banco Mundial) e "
            "livros de referência; 2000–2026; português, inglês ou espanhol; "
            "aderência a pelo menos um eixo temático.")
bullet(doc, "Exclusão: editoriais e cartas sem dado próprio; trabalhos sem "
            "método descrito; duplicatas; textos sem acesso ao conteúdo completo.")
bullet(doc, "Registrar o fluxo (identificados → após remoção de duplicatas → "
            "triados por título/resumo → lidos na íntegra → incluídos) numa "
            "tabela no estilo PRISMA, mesmo sem revisão sistemática formal.")

doc.add_heading("4.3 Modelo de fichamento (uma linha por artigo)", 2)
grid_table(doc,
    ["Campo", "O que preencher"],
    [
     ["Referência (ABNT NBR 6023)", "sobrenome, iniciais; título; periódico; v., n., p.; ano; DOI/URL"],
     ["Tipo de estudo", "revisão, ecológico, coorte, análise documental, relatório, ensaio teórico"],
     ["Objetivo do estudo", "1–2 frases"],
     ["Método / dados", "população, fontes, período, técnica analítica"],
     ["Principais achados", "3–5 pontos, com números quando houver"],
     ["Conceito/definição que empresta ao trabalho", "ex.: definição de DRSAI; de equidade horizontal/vertical; de needs-based allocation"],
     ["Como se relaciona com o Farol-SS", "sustenta a introdução? o método? a discussão? contrapõe algo?"],
     ["Citação-chave (trecho literal + página)", "1–2 trechos entre aspas, para citação direta"],
     ["Onde entra no texto", "Introdução / Metodologia / Resultados-discussão / Limitações"],
    ],
    [5.5, 10.8])
p(doc, "Sugestão operacional: manter os fichamentos numa planilha (uma aba por "
       "eixo) e exportar a bibliografia final para um gerenciador (Zotero/Mendeley) "
       "com o estilo ABNT.", size=9, color=GREY)

doc.add_heading("4.4 Bibliografia inicial — referências candidatas (CONFERIR dados completos)", 2)
p(doc, "Lista de partida por tema. Todas são obras reais e conhecidas na área, mas "
       "os dados bibliográficos completos (volume, página, ano exato, DOI) devem "
       "ser confirmados na base antes de citar. Não citar nada que não tenha sido "
       "lido e fichado.", italic=True, color=GREY)
for tema, refs in [
    ("Sistema de saúde brasileiro e financiamento", [
        "PAIM, J. et al. The Brazilian health system: history, advances, and challenges. The Lancet, 2011.",
        "CASTRO, M. C. et al. Brazil's unified health system: the first 30 years and prospects for the future. The Lancet, 2019.",
        "MENDES, A.; MARQUES, R. M. O financiamento do SUS sob os \"ventos\" da financeirização. Ciência & Saúde Coletiva.",
        "VIEIRA, F. S.; BENEVIDES, R. P. S. Os impactos do Novo Regime Fiscal para o financiamento do SUS. Ipea, Nota Técnica.",
        "PIOLA, S. F. et al. Financiamento público da saúde: uma história à procura de rumo. Ipea, Texto para Discussão.",
    ]),
    ("Equidade e alocação de recursos em saúde", [
        "WHITEHEAD, M. The concepts and principles of equity and health. International Journal of Health Services, 1992.",
        "RICE, N.; SMITH, P. C. Capitation and risk adjustment in health care financing. Milbank Quarterly.",
        "PORTO, S. M. et al. Metodologia de alocação equitativa de recursos / alocação de recursos financeiros no SUS (Fiocruz/DAD).",
        "ASTHANA, S.; GIBSON, A. Health care equity, health equity and resource allocation. (territorial/needs-based funding)",
    ]),
    ("Saneamento e saúde (DRSAI)", [
        "FUNASA. Manual de saneamento / classificação das Doenças Relacionadas ao Saneamento Ambiental Inadequado (DRSAI).",
        "HELLER, L. Saneamento e saúde. Brasília: OPAS/OMS, 1997.",
        "TEIXEIRA, J. C.; HELLER, L. e cols. Fatores ambientais associados às doenças diarreicas / associação saneamento–saúde.",
        "BARRETO, M. L. et al. Impact of a citywide sanitation program in Northeast Brazil on diarrhoea. The Lancet, 2007.",
    ]),
    ("Sistemas de informação e vigilância em saúde", [
        "BRASIL. Ministério da Saúde. SINAN: normas e rotinas / Sistema de Informação de Agravos de Notificação.",
        "BRASIL. Ministério da Saúde. SIH-SUS: manual técnico do Sistema de Informações Hospitalares.",
        "LIMA, C. R. A. et al. Qualidade dos dados dos sistemas de informação em saúde no Brasil (revisão).",
    ]),
    ("Dados abertos governamentais e reúso", [
        "JANSSEN, M.; CHARALABIDIS, Y.; ZUIDERWIJK, A. Benefits, adoption barriers and myths of open data and open government. Information Systems Management, 2012.",
        "ATTARD, J. et al. A systematic review of open government data initiatives. Government Information Quarterly, 2015.",
        "OPEN KNOWLEDGE / W3C. Open Data Handbook; Data on the Web Best Practices.",
        "BRASIL. CGU. Portal Brasileiro de Dados Abertos — documentação e diretrizes de reúso.",
    ]),
    ("Indicadores compostos e análise espacial em saúde", [
        "OECD; JRC. Handbook on Constructing Composite Indicators: methodology and user guide. 2008.",
        "BARCELLOS, C.; SANTOS, S. M. Colocando dados no mapa: a escolha da unidade espacial de agregação. Cad. Saúde Pública.",
        "CARVALHO, M. S.; SOUZA-SANTOS, R. Análise de dados espaciais em saúde pública. Cad. Saúde Pública.",
    ]),
    ("Integridade em contratações públicas de saúde", [
        "FAZEKAS, M.; TÓTH, I. J. Corruption risk indicators in public procurement (objective/red-flag approach).",
        "OPEN CONTRACTING PARTNERSHP. Red flags for integrity in public contracting.",
        "BRASIL. TCU / CGU. Referenciais de preços e sobrepreço em compras públicas de medicamentos e insumos.",
    ]),
]:
    doc.add_heading(tema, 3)
    for r in refs:
        bullet(doc, r)

# 5
doc.add_heading("5. Metodologia do Farol-SS (síntese para o texto)", 1)
if DIAG.exists():
    fig = doc.add_paragraph()
    fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig.add_run().add_picture(str(DIAG), width=Cm(12))
    p(doc, "Figura 1. Construção do IEAS — das oito fontes federais aos dois eixos "
           "ponderados, ao ranque percentil, ao gap e ao semáforo. (candidata a "
           "figura do pôster)", size=9, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
bullet(doc, "Recorte: 185 municípios de Pernambuco, 2020–2024; grão canônico "
            "município × ano (grade completa 185 × 5 = 925).")
bullet(doc, "Oito fontes federais abertas (todas catalogadas no dados.gov.br): "
            "SINAN (arboviroses e doenças de veiculação hídrica); SIH-SUS grupo RD "
            "(internações por DRSAI, por município de residência); Censo 2022 do "
            "IBGE (déficit de água, esgoto e coleta de lixo); CadÚnico/SAGI (taxa "
            "de famílias em extrema pobreza); SIOPS (despesa própria municipal em "
            "saúde por habitante); PNCP e Compras.gov.br (contratação de insumos, "
            "L3 municipal + federal); Portal da Transparência (transferências "
            "federais, proxy de L1).")
bullet(doc, "Pipeline reprodutível em camadas (bronze → silver → gold), DuckDB "
            "sobre Parquet, sem servidor de banco; proveniência (URL, SHA-256, "
            "contagem de linhas, data) registrada em manifest.json.")
bullet(doc, "Eixo Necessidade (N): média ponderada de três subíndices — "
            "epidemiológico 0,40 (arboviroses + veiculação hídrica + internações "
            "DRSAI), saneamento 0,35 (déficit ponderado água 0,35 / esgoto 0,45 / "
            "lixo 0,20), vulnerabilidade 0,25 (extrema pobreza).")
bullet(doc, "Eixo Alocação (A): média ponderada de três camadas per capita "
            "deflacionadas para 2024 pelo IPCA — L1 repasse federal 0,35, L2 "
            "execução própria 0,40, L3 contratação de insumos 0,25.")
bullet(doc, "Normalização: cada eixo é convertido em ranque percentil ∈ [0, 1] "
            "entre os 185 municípios (posição relativa, não valor absoluto — "
            "robustez a inflação, escala e outliers).")
bullet(doc, "Índice: gap = ranque(A) − ranque(N) ∈ [−1, 1] colore um semáforo de "
            "quatro cores; ieas = 1 − |gap| mede alinhamento. \"Regra do cinza\": "
            "sem cobertura mínima de componentes (60% em N, 50% em A) o índice "
            "não é publicado.")
bullet(doc, "Quatro detectores de anomalia com explicação em linguagem natural: "
            "(1) desalinhamento estrutural (farol vermelho); (2) alocação abaixo "
            "do esperado (resíduo robusto de regressão necessidade→alocação, por "
            "ano); (3) suspeita de sobrepreço (preço unitário acima de Q3 + 1,5·IQR "
            "da mesma categoria/unidade/dose); (4) suspeita de desabastecimento "
            "(incidência sustentada de um agravo sem contratação do insumo "
            "correspondente).")
bullet(doc, "Software e reprodutibilidade: ~6.100 linhas de Python; 75 testes "
            "automatizados; código sob domínio público.")

# 6
doc.add_heading("6. Resultados principais (síntese para o texto)", 1)
grid_table(doc,
    ["Indicador", "Valor"],
    [
     ["Município-anos com IEAS calculado", "921 de 925 (185/185 em 2024)"],
     ["Farol — todos os anos", "383 verde · 202 vermelho · 176 azul · 160 amarelo · 4 cinza"],
     ["Farol — 2024", "104 verde · 60 azul · 19 amarelo · 2 vermelho · 0 cinza"],
     ["Alertas emitidos", "777 (570 suspeita de desabastecimento · 202 desalinhamento "
      "estrutural · 4 alocação abaixo do esperado · 1 suspeita de sobrepreço)"],
     ["Volume ingerido", "271.505 notificações (SINAN) · 22.597 internações DRSAI (SIH) · "
      "6.150 contratos + 4.691 itens (PNCP)"],
     ["Cobertura mediana de saneamento em PE (Censo 2022)", "água 64% · esgoto 48% · coleta de lixo 74%"],
    ],
    [6.0, 10.3])
bullet(doc, "A distribuição do farol muda ano a ano: 2020–2022 concentram os "
            "vermelhos (necessidade no topo, alocação no fundo); 2023–2024 pendem "
            "para verde/azul — parte por crescimento real de repasse e execução, "
            "parte por sensibilidade do eixo Alocação ao peso de L1 (proxy).")
bullet(doc, "As baixas coberturas censitárias de saneamento sustentam necessidade "
            "elevada em quase todo o estado; o componente de internações DRSAI "
            "reforça o sinal onde a carga hospitalar por doença de veiculação "
            "hídrica é alta.")
bullet(doc, "Os alertas de desabastecimento concentram-se nos anos de surto de "
            "arbovirose (2022 e 2024).")

# 7
doc.add_heading("7. Rascunho do resumo (parágrafo único — ajustar para ≤ 2.500 caracteres)", 1)
resumo = (
 "A alocação de recursos públicos em saúde e a carga de doença não se distribuem "
 "de forma equivalente no território, e os portais de transparência informam "
 "quanto se gasta, mas não se o gasto acompanha a necessidade de cada município. "
 "O objetivo foi desenvolver e disponibilizar o Farol da Saúde & Saneamento "
 "(Farol-SS), monitor territorial de código aberto que quantifica o alinhamento "
 "entre necessidade sanitária e alocação de recursos nos 185 municípios de "
 "Pernambuco (2020–2024). Metodologia: um pipeline de dados reprodutível "
 "(arquitetura em camadas, DuckDB sobre Parquet, proveniência rastreável) "
 "integrou oito fontes federais abertas — notificações do SINAN, internações do "
 "SIH-SUS por doenças relacionadas ao saneamento, déficit de saneamento do Censo "
 "2022, extrema pobreza do CadÚnico, execução própria municipal em saúde "
 "(SIOPS), contratação de insumos (PNCP e Compras.gov.br) e transferências "
 "federais (Portal da Transparência) — num grão único de município × ano. "
 "Construiu-se o Índice de Efetividade da Alocação Sanitária (IEAS): dois eixos, "
 "Necessidade (epidemiológico, saneamento e vulnerabilidade social) e Alocação "
 "(repasse federal, execução própria e contratação de insumos, em reais per "
 "capita deflacionados), convertidos em ranque percentil dentro do estado; a "
 "diferença gap = ranque(Alocação) − ranque(Necessidade) colore um semáforo de "
 "quatro cores, e uma regra do cinza impede o cálculo quando a cobertura de "
 "dados é insuficiente. Quatro detectores geram alertas explicáveis em linguagem "
 "natural (desalinhamento estrutural, alocação abaixo do esperado, suspeita de "
 "sobrepreço e suspeita de desabastecimento de insumos). Resultados: o IEAS foi "
 "calculado para 921 dos 925 município-anos (185 de 185 em 2024); 202 "
 "município-anos ficaram em situação de necessidade não atendida e foram "
 "emitidos 777 alertas; as coberturas censitárias medianas de água (64%), esgoto "
 "(48%) e coleta de lixo (74%) sustentam necessidade elevada em quase todo o "
 "estado. A entrega inclui painel web, API aberta em JSON/CSV e relatório "
 "técnico, com proveniência de cada dado. Conclusões: é viável, a partir "
 "exclusivamente de dados abertos, produzir um instrumento reprodutível e "
 "explicável de apoio ao controle social e à priorização de auditoria e de "
 "investimento em saúde e saneamento, transferível a outras unidades da "
 "federação."
)
pr = doc.add_paragraph(resumo)
pr.paragraph_format.line_spacing = 1.3
p(doc, f"[Contagem automática: {len(resumo)} caracteres com espaços. "
       f"Limite do evento: 2.500. {'Dentro do limite.' if len(resumo) <= 2500 else 'REDUZIR.'} "
       f"Reconferir no formulário — a contagem pode diferir por conta de travessão, "
       f"× e sinal de menos.]",
  size=9, color=(AQUA if len(resumo) <= 2500 else RGBColor(0xC6, 0x28, 0x28)))
bullet(doc, "Se a contagem do formulário estourar 2.500: substituir a lista de "
            "fontes por \"oito fontes federais abertas\", enxugar a descrição dos "
            "quatro detectores e a frase sobre a entrega/proveniência.")
bullet(doc, "Revisar: uso de travessão (—) vs. hífen; \"ranque\" vs. \"rank\"; "
            "manter um único tempo verbal por seção.")

# 8
doc.add_heading("8. Estrutura do pôster (banner 120 × 90 cm)", 1)
p(doc, "Ver o arquivo banner-sipes-2026-modelo.docx (modelo em escala, já com a "
       "logo do evento). Conteúdo por bloco:", bold=True)
grid_table(doc,
    ["Bloco", "Conteúdo"],
    [
     ["Cabeçalho", "logo do V SIPES (obrigatória) · título · autores e filiações · e-mail do 1º autor"],
     ["Introdução", "o descompasso necessidade × alocação; lacuna que os portais de transparência não cobrem"],
     ["Objetivos", "geral + 2–3 específicos"],
     ["Metodologia", "Figura 1 (diagrama do IEAS) + texto curto: fontes, eixos, ranque percentil, regra do cinza, detectores"],
     ["Resultados", "mapa do farol de PE (2024) + tabela-resumo (921/925; distribuição do farol; 777 alertas) + 1 exemplo de alerta explicável"],
     ["Conclusões", "viabilidade com dados abertos; usos (controle social, auditoria, priorização); transferibilidade"],
     ["Limitações", "proxy de L1; saneamento = retrato 2022; maturidade do PNCP"],
     ["Rodapé", "referências (5–8, fonte reduzida) · QR do painel · \"V SIPES 2026 · Recife-PE\""],
    ],
    [3.5, 12.8])
bullet(doc, "Figuras a preparar: (a) diagrama do IEAS — já em assets/diagrama-ieas.png; "
            "(b) mapa coroplético do farol para 2024 — exportar da página Farol do "
            "painel; (c) opcional: recorte da página Alertas com um alerta em "
            "linguagem natural.")

# 9
doc.add_heading("9. Cronograma", 1)
grid_table(doc,
    ["Semana / data", "Entrega"],
    [
     ["Até 15/06/2026", "Definir autores, área temática e modalidade; abrir a planilha de fichamento"],
     ["16–30/06/2026", "Buscas nas bases; triagem por título/resumo; selecionar 15–20 textos"],
     ["01–15/07/2026", "Leitura na íntegra e fichamento (mínimo 12 fichas); fechar bibliografia"],
     ["16–24/07/2026", "Redigir o resumo (parágrafo único ≤ 2.500 caracteres); revisão de português; conferir palavras-chave no DeCS"],
     ["25–29/07/2026", "Revisão final e SUBMISSÃO na 1ª rodada (prazo 30/07/2026; resultado até 30/08)"],
     ["Contingência", "2ª rodada 30/08 (resultado 30/09) · 3ª rodada 30/09 (resultado 20/10)"],
     ["Após aceite", "Montar o banner no modelo; preparar mapa e QR; ensaiar apresentação (10+5 min se oral)"],
     ["25–27/11/2026", "Apresentação no evento (Recife-PE)"],
    ],
    [4.6, 11.7])

# 10
doc.add_heading("10. Divisão de tarefas", 1)
grid_table(doc,
    ["Frente", "Responsável", "Produto"],
    [
     ["Revisão de literatura e fichamento", "[ ]", "planilha de fichas + bibliografia ABNT"],
     ["Redação do resumo", "[ ]", "resumo final ≤ 2.500 caracteres"],
     ["Figuras (diagrama, mapa, alerta)", "[ ]", "PNGs em alta resolução"],
     ["Banner", "[ ]", "PDF 90 × 120 cm, 300 dpi, com logo do evento"],
     ["Submissão e acompanhamento", "[ ]", "comprovante + e-mails do comitê"],
     ["Apresentação", "[ ]", "roteiro de 10 min + respostas a perguntas prováveis"],
    ],
    [5.5, 3.5, 7.3])

# 11
doc.add_heading("11. Checklist de submissão", 1)
for item in [
    "Resumo em parágrafo único, sem seções, com introdução, objetivos, metodologia, resultados e conclusões",
    "Resumo com no máximo 2.500 caracteres COM espaços (conferir com acentuação)",
    "Título com no máximo 50 palavras, só a primeira letra maiúscula",
    "Palavras-chave: Saneamento [D012499]; Alocação de Recursos [D040841]; Gastos em Saúde [D005102] — confirmadas no DeCS",
    "No máximo 6 autores; 1º autor não é principal em outro trabalho; nome completo + e-mail + instituição de cada um",
    "Trabalho inédito (não apresentado em outro evento)",
    "Submissão feita pelo formulário on-line (não por e-mail)",
    "Área temática selecionada entre os 11 eixos do evento",
    "Revisão de português feita por pessoa diferente de quem redigiu",
    "Guardar comprovante de submissão e a data da rodada",
    "Para o pôster: banner 120 × 90 cm com a logomarca do V SIPES; PDF final em 300 dpi",
    "Reconferir todas as regras em sipes.com.br/trabalhos antes de enviar (podem ter mudado)",
]:
    par = doc.add_paragraph(item, style="List Bullet")
    par.runs[0].font.size = Pt(10.5)
    par.paragraph_format.space_after = Pt(2)

# 12
doc.add_heading("12. Riscos e pendências", 1)
grid_table(doc,
    ["Risco / pendência", "Mitigação"],
    [
     ["Normas do site incompletas (orientação/arquivo do pôster, citação)",
      "Escrever para contato@sipes.com.br pedindo o modelo oficial e a especificação do e-pôster"],
     ["\"Ineditismo\": o Farol-SS foi submetido ao Concurso da CGU",
      "Concurso não é evento científico com anais; ainda assim, declarar no formulário e, se houver dúvida, consultar o comitê"],
     ["Resumo estourar 2.500 caracteres", "Ter a versão curta pronta (seção 7)"],
     ["Aderência temática", "Enquadrar em \"inovação tecnológica em saúde\"; reforçar o elo com vigilância e determinantes (saneamento) — eixos do evento"],
     ["Dados do painel podem ser atualizados", "Congelar a versão citada (commit e data) e informar no texto/pôster"],
    ],
    [7.0, 9.3])

doc.add_paragraph()
p(doc, "Documento gerado a partir do estado do projeto Farol-SS em 30/08/2026. "
       "Regenerar com docs/sipes-2026/_build_projeto.py.", size=8, color=GREY,
  align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(str(OUT))
print("ok:", OUT, OUT.stat().st_size, "bytes")
