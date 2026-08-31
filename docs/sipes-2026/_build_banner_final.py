#!/usr/bin/env python3
"""Gera docs/sipes-2026/banner-final-sipes-2026.docx — versão final do pôster, com o
texto argumentativo, as citações da revisão aplicadas em sequência (NBR
10520:2023), gráficos, tabela de resultados e as referências (NBR 6023:2018).

Escala 1:3 de 90 x 120 cm (retrato). Rodar antes: 03, 05 e _graficos.py.
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AQUA, AQUA_LIGHT = "1C8CB0", "E8F4F8"
NAVY = RGBColor(0x12, 0x30, 0x3F)
GREY = RGBColor(0x53, 0x5C, 0x6B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CIT = RGBColor(0x0D, 0x45, 0x85)   # citações — azul institucional

HERE = Path(__file__).parent
A = HERE / "assets"
RIS = HERE / "revisao" / "referencias.ris"
DAT = json.loads((HERE / "revisao" / "_farol_dados.json").read_text(encoding="utf-8"))
OUT = HERE / "banner-final-sipes-2026.docx"

CITADOS = ["barcellos_2005", "carvalho_2021", "chaves_2017", "cohn_2005",
           "correia_2014", "duarte_2019", "funcia_2019", "lima_2009", "luiz_2009",
           "massuda_2018", "mendes_2011", "neves-silva_2016", "paiva_2018",
           "rodrigues_2021", "santos-neto_2017", "silva_2009", "viacava_2019"]


def refs_abnt():
    txt = RIS.read_text(encoding="utf-8", errors="ignore")
    blocos = txt.split("ER  - ")
    out = []
    for b in blocos:
        m = re.search(r"referência ABNT:\s*(.+)", b)
        ch = re.search(r"chave:\s*(\S+)", b)
        if m and ch and any(ch.group(1).startswith(c) for c in CITADOS):
            r = re.sub(r"<[^>]+>", "", m.group(1)).strip()
            out.append(r)
    return sorted(out)


def shade(el, fill):
    pr = el._tc.get_or_add_tcPr() if hasattr(el, "_tc") else el._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill)
    pr.append(s)


def no_borders(t):
    bd = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}")
        x.set(qn("w:val"), "none")
        bd.append(x)
    t._tbl.tblPr.append(bd)


def para(cell, runs, *, align=None, after=4, before=0, sp=1.12):
    p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = sp
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, sty in runs:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(sty.get("size", 9.3))
        r.bold = sty.get("bold", False)
        r.italic = sty.get("italic", False)
        r.font.color.rgb = sty.get("color", NAVY)
    return p


def c(txt):
    return (txt, {"color": CIT, "bold": True, "size": 9.3})


def T(txt):
    return (txt, {})


def sec(cell, text):
    p = para(cell, [(text, {"size": 12, "bold": True, "color": WHITE})], before=6, after=3)
    shade(p, AQUA)
    p.paragraph_format.left_indent = Pt(4)


def just(cell, *runs):
    return para(cell, list(runs), align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def fig(cell, nome, legenda, largura=12.4):
    fp = cell.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(3)
    fp.paragraph_format.space_after = Pt(1)
    if (A / nome).exists():
        fp.add_run().add_picture(str(A / nome), width=Cm(largura))
    para(cell, [(legenda, {"size": 7.6, "color": GREY})],
         align=WD_ALIGN_PARAGRAPH.CENTER, after=4)


# ── documento ───────────────────────────────────────────────────────
doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Cm(30), Cm(40)
s.left_margin = s.right_margin = Cm(1.3)
s.top_margin = s.bottom_margin = Cm(1.1)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(9.3)

para(doc, [("Versão final do pôster · escala 1:3 (final 90 × 120 cm, retrato, "
            "300 dpi) · citações em ", {"size": 7.3, "color": GREY, "italic": True}),
           ("azul", {"size": 7.3, "color": CIT, "italic": True, "bold": True}),
           (" = da revisão (NBR 10520:2023) · logo do V SIPES obrigatória",
            {"size": 7.3, "color": GREY, "italic": True})],
     align=WD_ALIGN_PARAGRAPH.CENTER, after=5)

# cabeçalho
h = doc.add_table(rows=1, cols=2)
no_borders(h)
h.alignment = WD_TABLE_ALIGNMENT.CENTER
h.columns[0].width, h.columns[1].width = Cm(9), Cm(18.4)
hc0, hc1 = h.rows[0].cells
shade(hc0, AQUA_LIGHT)
shade(hc1, AQUA_LIGHT)
if (A / "sipes-logo.png").exists():
    hc0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    hc0.paragraphs[0].add_run().add_picture(str(A / "sipes-logo.png"), width=Cm(8.2))
para(hc1, [("V Simpósio Internacional de Pesquisa em Estilos de Vida & Saúde",
            {"size": 10.5, "bold": True})], align=WD_ALIGN_PARAGRAPH.RIGHT, after=2)
para(hc1, [("Recife-PE · 25 a 27 de novembro de 2026 · Modalidade: Pôster",
            {"size": 9, "color": GREY})], align=WD_ALIGN_PARAGRAPH.RIGHT)

para(doc, "", after=1)
para(doc, [("Farol da Saúde & Saneamento: um índice territorial de efetividade "
            "da alocação sanitária construído com dados abertos para os 185 "
            "municípios de Pernambuco, 2020–2024", {"size": 16, "bold": True})],
     align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
para(doc, [("SOBRENOME, Nome¹; SOBRENOME, Nome²  [até 6 autores]",
            {"size": 10})], align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
para(doc, [("¹Instituição, cidade/UF · ²Instituição · e-mail do autor apresentador",
            {"size": 8, "color": GREY})], align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
para(doc, [("Palavras-chave (DeCS): ", {"size": 8.5, "bold": True, "color": GREY}),
           ("Saneamento; Alocação de Recursos; Gastos em Saúde",
            {"size": 8.5, "color": GREY})], align=WD_ALIGN_PARAGRAPH.CENTER, after=5)

# corpo — 2 colunas
b = doc.add_table(rows=1, cols=2)
no_borders(b)
b.alignment = WD_TABLE_ALIGNMENT.CENTER
b.columns[0].width = b.columns[1].width = Cm(13.6)
L, R = b.rows[0].cells

# ── coluna esquerda: introdução, objetivos, metodologia ─────────────
sec(L, "1  INTRODUÇÃO")
just(L,
     T("A alocação de recursos públicos em saúde e a carga de doença não se "
       "distribuem de forma equivalente no território brasileiro; a alocação no "
       "SUS é descrita como subótima e como origem de disparidades regionais de "
       "acesso e de desfechos "), c("(MASSUDA et al., 2018)"),
     T(". Apesar das exigências legais de universalidade e equidade, essa "
       "alocação segue, em regra, o comportamento histórico pautado na produção "
       "de serviços, e não nas necessidades da população "),
     c("(MENDES et al., 2011)"),
     T("; os critérios legais de partilha, pensados para reduzir a desigualdade "
       "regional, ainda não foram efetivamente implementados "),
     c("(CARVALHO, 2021)"),
     T(", e o teto de gastos comprime o recurso por habitante ao longo do tempo "),
     c("(FUNCIA, 2019)"),
     T(". O déficit de saneamento, por sua vez, tem efeito mensurável: cerca de "
       "16% das internações por doenças de veiculação hídrica no país seriam "
       "evitáveis com esgotamento adequado "), c("(PAIVA; SOUZA, 2018)"),
     T(" — ainda que a magnitude atribuível ao saneamento seja menor do que o "
       "senso comum sugere e exija rigor na definição do indicador "),
     c("(BARCELLOS, 2005)"),
     T(". Os portais de transparência informam quanto se gasta, mas não "
       "permitem avaliar se o gasto acompanha a necessidade de cada município."))

sec(L, "2  OBJETIVOS")
just(L, T("Desenvolver e disponibilizar o Farol da Saúde & Saneamento (Farol-SS), "
          "monitor territorial de código aberto que quantifica o alinhamento "
          "entre necessidade sanitária e alocação de recursos nos 185 municípios "
          "de Pernambuco (2020–2024), e descrever o Índice de Efetividade da "
          "Alocação Sanitária (IEAS) e os alertas dele derivados."))

sec(L, "3  METODOLOGIA")
just(L,
     T("Estudo ecológico de base documental, com dados secundários de acesso "
       "aberto. Um pipeline reprodutível (camadas bronze–silver–gold; DuckDB "
       "sobre arquivos Parquet; proveniência com SHA-256) integrou oito fontes "
       "federais num grão único de município × ano (grade 185 × 5) — mesmo "
       "princípio de reúso e cruzamento de bases que revela lacunas invisíveis "
       "a uma fonte isolada "), c("(SILVA; LEITE; ALMEIDA, 2009)"), T("."))
just(L,
     ("Fórmula. ", {"bold": True, "size": 9.3}),
     T("O IEAS compara dois eixos. Cada eixo é a média ponderada de seus "
       "subíndices, com pesos fixados em arquivo de configuração versionado — "
       "Necessidade: epidemiológico 0,40; saneamento 0,35; vulnerabilidade "
       "0,25; Alocação: repasse federal 0,35; execução própria 0,40; "
       "contratação de insumos 0,25, em R$/hab deflacionados pelo IPCA. Cada "
       "média é convertida em ranque percentil ∈ [0, 1] entre os 185 "
       "municípios: N_m é a posição relativa da necessidade e A_m a da "
       "alocação. O índice é gap_m = A_m − N_m ∈ [−1, 1] (gap ≤ −0,33 = "
       "necessidade não atendida), e ieas_m = 1 − |gap_m|, usado apenas para "
       "ordenar."))
just(L,
     ("Escolhas estatísticas. ", {"bold": True, "size": 9.3}),
     T("Índices sintéticos de condições de vida para municípios brasileiros "
       "costumam padronizar por escore-z e agrupar por análise de conglomerados "),
     c("(LUIZ et al., 2009)"),
     T(", ou compor um índice de necessidades por análise de componentes "
       "principais "), c("(MENDES et al., 2011)"),
     T(". O Farol-SS usa o ranque percentil: limitado a [0, 1] e independente "
       "de distribuição, não pressupõe normalidade e impede que um valor "
       "extremo de um município pequeno distorça a escala dos demais; e uma "
       "média ponderada explícita, que troca parcimônia estatística por "
       "transparência e reprodutibilidade. Componente ausente numa linha tem o "
       "peso redistribuído entre os presentes; abaixo de 60% (Necessidade) ou "
       "50% (Alocação) de componentes presentes o índice não é publicado — a "
       "completitude dos sistemas de informação em saúde é heterogênea e sem "
       "método padronizado "),
     c("(CORREIA; PADILHA; VASCONCELOS, 2014; LIMA et al., 2009)"), T("."))
just(L,
     ("Detectores. ", {"bold": True, "size": 9.3}),
     T("Quatro rotinas geram alertas explicáveis: (i) todo farol vermelho; "
       "(ii) resíduo de uma regressão anual necessidade→alocação padronizado "
       "por 1,4826·MAD (desvio absoluto mediano), resistente aos próprios "
       "outliers que se quer detectar; (iii) preço unitário de insumo acima de "
       "Q3 + 1,5·IQR da mesma categoria, unidade e dose no estado (cerca de "
       "Tukey); (iv) incidência de agravo no percentil 75+ do estado sem "
       "contratação do insumo correspondente. O SIOPS fornece a despesa própria "
       "per capita — que varia em ordem de grandeza entre municípios vizinhos "),
     c("(SANTOS NETO et al., 2017)"),
     T(" —; o PNCP fornece o preço unitário, num campo em que o poder de compra "
       "do Estado é decisivo "),
     c("(CHAVES; OSORIO-DE-CASTRO; OLIVEIRA, 2017)"), T("."))
fig(L, "diagrama-ieas.png", "Figura 1 – Construção do IEAS. Fonte: os autores (2026).",
    largura=11.6)

# ── coluna direita: resultados, conclusão ──────────────────────────
sec(R, "4  RESULTADOS")
just(R,
     T("O IEAS foi calculado para 921 dos 925 município-anos (185 de 185 em "
       "2024). A distribuição do semáforo mostra 383 município-anos alinhados, "
       "202 em necessidade não atendida, 176 com alocação acima da necessidade "
       "e 160 com subalocação leve (Figura 2)."))
fig(R, "fig1_farol_dist.png", "Figura 2 – Distribuição do farol, PE, 2020–2024. Fonte: Farol-SS (2026).")
just(R,
     T("A leitura muda ao longo do período: 2020–2022 concentram os municípios "
       "em necessidade não atendida e 2023–2024 pendem para verde e azul "
       "(Figura 3) — virada em parte real (repasse e execução crescem) e em "
       "parte sensível ao peso da camada de repasse federal (proxy)."))
fig(R, "fig2_farol_ano.png", "Figura 3 – Farol por ano. Fonte: Farol-SS (2026).")
just(R,
     T("Em 2024, a maior parte dos municípios está sobre ou acima da diagonal "
       "necessidade = alocação; os pontos abaixo dela são os faróis amarelo e "
       "vermelho (Figura 4)."))
fig(R, "fig3_scatter_2024.png", "Figura 4 – Necessidade × alocação (ranque percentil), PE, 2024. Fonte: Farol-SS (2026).",
    largura=9.6)
just(R,
     T("Foram emitidos 777 alertas explicáveis; 570 de suspeita de "
       "desabastecimento de insumos, concentrados nos anos de surto de "
       "arbovirose (Figura 5) — padrão coerente com a sensibilidade da carga "
       "de diarreia a fatores ambientais e climáticos "),
     c("(DUARTE et al., 2019)"), T("."))
fig(R, "fig4_desabastecimento.png", "Figura 5 – Alertas de desabastecimento por ano. Fonte: Farol-SS (2026).")

# tabela de resultados
para(R, [("Tabela 1 – Síntese dos resultados (2020–2024)", {"size": 8.5, "bold": True})],
     before=4, after=2)
tb = R.add_table(rows=0, cols=2)
tb.style = "Light Grid Accent 1"
for kk, vv in [
    ("Município-anos com IEAS calculado", "921 de 925 (185/185 em 2024)"),
    ("Farol — total (todos os anos)", "383 verde · 202 vermelho · 176 azul · 160 amarelo · 4 cinza"),
    ("Farol — 2024", "104 verde · 60 azul · 19 amarelo · 2 vermelho"),
    ("Alertas emitidos", "777 (570 desabastecimento · 202 desalinhamento · 4 resíduo · 1 sobrepreço)"),
    ("Cobertura mediana de saneamento (Censo 2022)", "água 64% · esgoto 48% · coleta de lixo 74%"),
]:
    r = tb.add_row().cells
    r[0].text = kk
    r[1].text = vv
    for cc in r:
        for pp in cc.paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(8)
    r[0].paragraphs[0].runs[0].bold = True
for row in tb.rows:
    row.cells[0].width = Cm(4.6)
    row.cells[1].width = Cm(8.8)
para(R, "", after=3)

sec(R, "5  CONCLUSÃO")
just(R,
     T("É viável, a partir exclusivamente de dados abertos, produzir um "
       "instrumento reprodutível e explicável de monitoramento territorial — "
       "resposta à necessidade reconhecida de acompanhar as desigualdades em "
       "saúde por região e grupo social "), c("(VIACAVA et al., 2019)"),
     T(" sem novo inquérito. O índice atua no plano da dotação de recursos "
       "frente à necessidade, um dos planos da equidade no financiamento "),
     c("(COHN, 2005)"),
     T(", e funciona como ferramenta de transparência e controle social sobre "
       "um direito — o saneamento — que condiciona a saúde "),
     c("(NEVES-SILVA; HELLER, 2016)"),
     T(". Como a eficiência do gasto municipal não acompanha a riqueza "),
     c("(RODRIGUES et al., 2021)"),
     T(", olhar necessidade × alocação, e não valores absolutos, é o recorte "
       "mais informativo. Limitações: a camada de repasse federal é um proxy; "
       "o saneamento é um retrato de 2022; a adesão ao PNCP cresce ano a ano."))

para(doc, "", after=2)

# rodapé — referências + identificação
f = doc.add_table(rows=1, cols=2)
no_borders(f)
f.columns[0].width, f.columns[1].width = Cm(20.6), Cm(6.6)
fc0, fc1 = f.rows[0].cells
shade(fc0, AQUA_LIGHT)
shade(fc1, AQUA_LIGHT)
para(fc0, [("REFERÊNCIAS (ABNT NBR 6023:2018)", {"size": 8, "bold": True})], after=2)
for r in refs_abnt():
    pp = para(fc0, [(r, {"size": 6.5})], after=1, sp=1.0)
    pp.paragraph_format.left_indent = Pt(0)
para(fc1, [("Painel: farol-ss.streamlit.app", {"size": 7.5})],
     align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
para(fc1, [("[ QR do painel ]", {"size": 7.5, "italic": True,
                                 "color": RGBColor(0x9A, 0xA3, 0xB0)})],
     align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
para(fc1, [("Código e dados: github.com/protazoarium/farol-ss",
            {"size": 6.5, "color": GREY})], align=WD_ALIGN_PARAGRAPH.CENTER)

para(doc, [("V SIPES 2026 · Recife-PE · Dados sob domínio público das fontes "
            "federais (IBGE, DATASUS, PNCP, CGU, MDS).", {"size": 7, "color": GREY})],
     align=WD_ALIGN_PARAGRAPH.CENTER, before=3)

doc.save(str(OUT))
print("ok:", OUT, OUT.stat().st_size, "bytes ·", len(refs_abnt()), "referências")
