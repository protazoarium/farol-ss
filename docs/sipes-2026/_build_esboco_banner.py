#!/usr/bin/env python3
"""Gera docs/sipes-2026/esboco-banner-sipes-2026.docx — ESBOÇO do trabalho em
formato de banner/pôster, com o conteúdo do Farol-SS já redigido em rascunho,
seguindo a estrutura de pôster técnico-científico da ABNT NBR 15437:2006 e
citações no sistema autor-data (NBR 10520:2023); referências em NBR 6023:2018.

As marcas (AUTOR, ANO) são pontos onde entram as citações da revisão — trocar
pelas fichas de revisao/fichamento.ris depois de lidas. As referências no rodapé
são puxadas de revisao/referencias.ris quando existir; senão, ficam de exemplo.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AQUA = "1C8CB0"
AQUA_LIGHT = "E8F4F8"
NAVY = RGBColor(0x12, 0x30, 0x3F)
GREY = RGBColor(0x5A, 0x62, 0x70)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CITE = RGBColor(0xB0, 0x53, 0x10)  # marcas de citação a preencher

HERE = Path(__file__).parent
LOGO = HERE / "assets" / "sipes-logo.png"
DIAG = HERE / "assets" / "diagrama-ieas.png"
RIS = HERE / "revisao" / "referencias.ris"
OUT = HERE / "esboco-banner-sipes-2026.docx"


def shade(el, fill):
    pr = el._tc.get_or_add_tcPr() if hasattr(el, "_tc") else el._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), fill)
    pr.append(s)


def no_borders(t):
    b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}")
        x.set(qn("w:val"), "none")
        b.append(x)
    t._tbl.tblPr.append(b)


def add(cell, runs, *, align=None, after=4, before=0, sp=1.12):
    """runs: str ou lista de (texto, dict-de-estilo)."""
    p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = sp
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, sty in runs:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(sty.get("size", 9.5))
        r.bold = sty.get("bold", False)
        r.italic = sty.get("italic", False)
        r.font.color.rgb = sty.get("color", NAVY)
    return p


def cite(txt="(AUTOR, ANO)"):
    return (txt, {"color": CITE, "bold": True, "size": 9})


def sec_title(cell, text):
    p = add(cell, [(text, {"size": 12.5, "bold": True, "color": WHITE})],
            before=6, after=3)
    shade(p, AQUA)
    p.paragraph_format.left_indent = Pt(4)


def justify(cell, *runs):
    return add(cell, list(runs), align=WD_ALIGN_PARAGRAPH.JUSTIFY)


# ── referências da revisão (NBR 6023) ────────────────────────────────
def referencias_abnt(limite=8):
    if not RIS.exists():
        return [
            "[As referências virão de revisao/referencias.ris após rodar o "
            "workflow. Exemplo de formatação NBR 6023:2018:]",
            "SOBRENOME, N. N. Título do artigo. <i>Nome do Periódico</i>, v. 00, "
            "n. 0, p. 00-00, 2024. DOI: https://doi.org/10.0000/xxxx. Acesso em: "
            "15 mar. 2026.",
        ]
    txt = RIS.read_text(encoding="utf-8", errors="ignore")
    refs = re.findall(r"referência ABNT:\s*(.+)", txt)
    refs = [re.sub(r"<[^>]+>", "", r).strip() for r in refs]
    return refs[:limite] or ["[referencias.ris sem campo 'referência ABNT' — rodar 01 de novo]"]


# ── documento ───────────────────────────────────────────────────────
doc = Document()
s = doc.sections[0]
s.page_width, s.page_height = Cm(30), Cm(40)     # escala 1:3 de 90 × 120 cm
s.left_margin = s.right_margin = Cm(1.4)
s.top_margin = s.bottom_margin = Cm(1.2)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10)

add(doc, [("ESBOÇO — pôster técnico-científico (ABNT NBR 15437:2006). Escala 1:3; "
           "final 90 × 120 cm, retrato, 300 dpi. Marcas ", {"size": 7.5, "color": GREY, "italic": True}),
          ("(AUTOR, ANO)", {"size": 7.5, "color": CITE, "italic": True, "bold": True}),
          (" = trocar pela citação da revisão. Logo do V SIPES obrigatória.",
           {"size": 7.5, "color": GREY, "italic": True})],
    align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

# cabeçalho
h = doc.add_table(rows=1, cols=2)
no_borders(h)
h.alignment = WD_TABLE_ALIGNMENT.CENTER
h.columns[0].width, h.columns[1].width = Cm(9.5), Cm(17.7)
c0, c1 = h.rows[0].cells
shade(c0, AQUA_LIGHT)
shade(c1, AQUA_LIGHT)
if LOGO.exists():
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c0.paragraphs[0].add_run().add_picture(str(LOGO), width=Cm(8.6))
add(c1, [("V Simpósio Internacional de Pesquisa em Estilos de Vida & Saúde",
          {"size": 11, "bold": True})], align=WD_ALIGN_PARAGRAPH.RIGHT, after=2)
add(c1, [("Recife-PE · 25 a 27 de novembro de 2026", {"size": 9.5, "color": GREY})],
    align=WD_ALIGN_PARAGRAPH.RIGHT, after=2)
add(c1, [("Pôster · Eixo: Inovação tecnológica em saúde (a confirmar)",
          {"size": 9, "color": RGBColor(0x1C, 0x8C, 0xB0)})],
    align=WD_ALIGN_PARAGRAPH.RIGHT)

add(doc, "", after=2)
add(doc, [("Farol da Saúde & Saneamento: um índice territorial de efetividade da "
           "alocação sanitária construído com dados abertos para os 185 municípios "
           "de Pernambuco, 2020–2024", {"size": 17, "bold": True})],
    align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
add(doc, [("SOBRENOME, Nome¹; SOBRENOME, Nome²  [até 6 autores; ¹e ² = filiação]",
           {"size": 10.5})], align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
add(doc, [("¹Instituição, cidade/UF · ²Instituição · e-mail do autor apresentador",
           {"size": 8.5, "color": GREY})], align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
add(doc, [("Palavras-chave (DeCS): ", {"size": 9, "bold": True, "color": GREY}),
          ("Saneamento; Alocação de Recursos; Gastos em Saúde",
           {"size": 9, "color": GREY})],
    align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

# corpo 2 colunas
b = doc.add_table(rows=1, cols=2)
no_borders(b)
b.alignment = WD_TABLE_ALIGNMENT.CENTER
b.columns[0].width = b.columns[1].width = Cm(13.4)
L, R = b.rows[0].cells

sec_title(L, "1 INTRODUÇÃO")
justify(L, ("A alocação de recursos públicos em saúde e a carga de doença não se "
            "distribuem de forma equivalente no território brasileiro ", {}),
        cite("(AUTOR, ANO)"),
        (". Municípios com maior déficit de saneamento e maior incidência de "
         "doenças de veiculação hídrica podem, ao mesmo tempo, receber e executar "
         "menos recurso por habitante do que a média do estado ", {}),
        cite("(AUTOR, ANO)"),
        (". Os portais de transparência informam o quanto se gasta, mas não "
         "permitem avaliar se esse gasto acompanha a necessidade sanitária de "
         "cada território ", {}),
        cite("(AUTOR, ANO)"),
        (".", {}))
add(L, [("[inserir 1–2 frases de fechamento com a lacuna que o trabalho preenche]",
         {"size": 8.5, "italic": True, "color": GREY})])

sec_title(L, "2 OBJETIVOS")
justify(L, ("Desenvolver e disponibilizar o Farol da Saúde & Saneamento (Farol-SS), "
            "monitor territorial de código aberto que quantifica o alinhamento "
            "entre necessidade sanitária e alocação de recursos nos 185 municípios "
            "de Pernambuco (2020–2024); descrever a construção do Índice de "
            "Efetividade da Alocação Sanitária (IEAS); e apresentar o panorama "
            "estadual e os alertas gerados.", {}))

sec_title(L, "3 METODOLOGIA")
justify(L, ("Estudo ecológico, de base documental, com dados secundários de acesso "
            "aberto. Um pipeline reprodutível (arquitetura em camadas; DuckDB "
            "sobre arquivos Parquet; proveniência com SHA-256) integrou oito "
            "fontes federais num grão único de município × ano (grade 185 × 5): "
            "SINAN; internações do SIH-SUS por doenças relacionadas ao saneamento "
            "ambiental inadequado ", {}), cite("(AUTOR, ANO)"),
        ("; déficit de água, esgoto e coleta de lixo do Censo Demográfico 2022; "
         "extrema pobreza do CadÚnico; despesa própria municipal em saúde (SIOPS); "
         "contratação de insumos (PNCP e Compras.gov.br); e transferências "
         "federais (Portal da Transparência).", {}))
justify(L, ("O IEAS combina dois eixos — Necessidade (epidemiológico 0,40; "
            "saneamento 0,35; vulnerabilidade 0,25) e Alocação (repasse federal "
            "0,35; execução própria 0,40; contratação de insumos 0,25, em R$/hab "
            "deflacionados) —, cada um convertido em ranque percentil ∈ [0,1] "
            "entre os 185 municípios. A diferença gap = ranque(Alocação) − "
            "ranque(Necessidade) define um semáforo de quatro cores; uma regra de "
            "cobertura mínima impede o cálculo quando faltam dados. Quatro "
            "detectores geram alertas explicáveis em linguagem natural. "
            "Indicadores compostos seguem recomendações da literatura ", {}),
        cite("(AUTOR, ANO)"), (".", {}))
if DIAG.exists():
    fp = L.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(4)
    fp.add_run().add_picture(str(DIAG), width=Cm(10.4))
    add(L, [("Figura 1 – Construção do IEAS. Fonte: elaborada pelos autores (2026).",
             {"size": 7.5, "color": GREY})], align=WD_ALIGN_PARAGRAPH.CENTER)

sec_title(R, "4 RESULTADOS")
justify(R, ("O IEAS foi calculado para 921 dos 925 município-anos (185 de 185 em "
            "2024). Na série completa, o semáforo distribuiu-se em 383 município-"
            "anos alinhados, 202 com necessidade não atendida, 176 com alocação "
            "acima da necessidade, 160 com subalocação leve e 4 sem dado "
            "suficiente; em 2024, 104 verde, 60 azul, 19 amarelo e 2 vermelho.", {}))
justify(R, ("Foram emitidos 777 alertas: 570 de suspeita de desabastecimento de "
            "insumos (concentrados nos anos de surto de arbovirose), 202 de "
            "desalinhamento estrutural, 4 de alocação abaixo do esperado e 1 de "
            "suspeita de sobrepreço. As coberturas censitárias medianas de água "
            "(64%), esgoto (48%) e coleta de lixo (74%) em Pernambuco sustentam "
            "necessidade elevada em quase todo o estado ", {}), cite("(AUTOR, ANO)"),
        (".", {}))
fp2 = R.add_paragraph()
fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
rr = fp2.add_run("[  INSERIR: mapa coroplético do farol — Pernambuco, 2024  ]\n"
                 "[  exportar da página Farol do painel · PNG 300 dpi  ]")
rr.font.size = Pt(9)
rr.italic = True
rr.font.color.rgb = RGBColor(0x9A, 0xA3, 0xB0)
shade(fp2, "F1F3F6")
add(R, [("Figura 2 – Farol da alocação sanitária, Pernambuco, 2024. "
         "Fonte: Farol-SS (2026).", {"size": 7.5, "color": GREY})],
    align=WD_ALIGN_PARAGRAPH.CENTER)

sec_title(R, "5 CONCLUSÃO")
justify(R, ("É viável, a partir exclusivamente de dados abertos, produzir um "
            "instrumento reprodutível e explicável de apoio ao controle social e à "
            "priorização de auditoria e de investimento em saúde e saneamento. O "
            "método é transferível a outras unidades da federação. Como "
            "limitações, a camada de repasse federal é um proxy de transferências "
            "sociais; o saneamento é um retrato de 2022; e a adesão ao Portal "
            "Nacional de Contratações Públicas ainda cresce ano a ano.", {}))

add(doc, "", after=2)

# rodapé — referências + identificação
f = doc.add_table(rows=1, cols=2)
no_borders(f)
f.columns[0].width, f.columns[1].width = Cm(20.4), Cm(6.8)
fc0, fc1 = f.rows[0].cells
shade(fc0, AQUA_LIGHT)
shade(fc1, AQUA_LIGHT)
add(fc0, [("REFERÊNCIAS (ABNT NBR 6023:2018) — só as citadas no pôster",
           {"size": 8.5, "bold": True})], after=2)
for ref in referencias_abnt(8):
    ref = ref.replace("<i>", "").replace("</i>", "")
    p = add(fc0, [(ref, {"size": 7})], after=2, sp=1.0)
    p.paragraph_format.left_indent = Pt(0)
add(fc1, [("Painel: farol-ss.streamlit.app", {"size": 7.5})],
    align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
add(fc1, [("[ QR do painel ]", {"size": 7.5, "italic": True,
                                "color": RGBColor(0x9A, 0xA3, 0xB0)})],
    align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
add(fc1, [("github.com/protazoarium/farol-ss", {"size": 6.5, "color": GREY})],
    align=WD_ALIGN_PARAGRAPH.CENTER)

add(doc, [("V SIPES 2026 · Recife-PE · Dados sob domínio público das fontes "
           "federais (IBGE, DATASUS, PNCP, CGU, MDS). Apoio: [agência/instituição].",
           {"size": 7, "color": GREY})],
    align=WD_ALIGN_PARAGRAPH.CENTER, before=4)

doc.save(str(OUT))
print("ok:", OUT, OUT.stat().st_size, "bytes  ·  referencias.ris presente:", RIS.exists())
